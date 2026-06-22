"""Convert PIPELINE_DATA_GAPS.md to PIPELINE_DATA_GAPS.docx.

Markdown features supported: headings (#-####), bold (**), italic (*), inline
code (`), tables (GitHub flavor), bullet/numbered lists, horizontal rules,
code blocks (```), and plain paragraphs. Output is a clean Word file with
proper styles applied, not an HTML→DOCX round-trip.
"""
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


MD = Path(__file__).resolve().parents[1] / "PIPELINE_DATA_GAPS.md"
DOCX = MD.with_suffix(".docx")

INLINE_PATTERN = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")


def add_inline_runs(paragraph, text: str) -> None:
    """Split text by **bold** / *italic* / `code` markers and add runs."""
    parts = INLINE_PATTERN.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:
            paragraph.add_run(part)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    """Consume a GitHub-flavored markdown table starting at lines[start].
    Returns (rows, next_index). rows[0] is the header."""
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].lstrip().startswith("|"):
        line = lines[i].strip()
        cells = [c.strip() for c in line.strip("|").split("|")]
        # Skip the separator row like |---|---|
        if not all(set(c) <= set("-: ") for c in cells):
            rows.append(cells)
        i += 1
    return rows, i


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    # pad rows so docx is happy
    rows = [r + [""] * (ncols - len(r)) for r in rows]
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Light Grid Accent 1"
    for r, row in enumerate(rows):
        for c, cell in enumerate(row):
            tc = table.cell(r, c)
            tc.text = ""
            p = tc.paragraphs[0]
            add_inline_runs(p, cell)
            if r == 0:
                for run in p.runs:
                    run.bold = True
    doc.add_paragraph()  # spacer


def convert(md_text: str, doc: Document) -> None:
    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_lang = ""
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Code fences
        if stripped.startswith("```"):
            in_code = not in_code
            code_lang = stripped[3:].strip() if in_code else ""
            if not in_code:
                # closing — already emitted the lines below
                pass
            i += 1
            continue
        if in_code:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            i += 1
            continue

        # Horizontal rule
        if stripped in ("---", "***", "___"):
            p = doc.add_paragraph()
            run = p.add_run("─" * 80)
            run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            heading = doc.add_heading("", level=min(level, 4))
            add_inline_runs(heading, text)
            i += 1
            continue

        # Table
        if stripped.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue

        # Bullet list
        m = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if m:
            indent = len(m.group(1)) // 2
            text = m.group(2)
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.25 * (indent + 1))
            add_inline_runs(p, text)
            i += 1
            continue

        # Numbered list
        m = re.match(r"^(\s*)\d+\.\s+(.*)$", line)
        if m:
            text = m.group(2)
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, text)
            i += 1
            continue

        # Blank line
        if not stripped:
            i += 1
            continue

        # Plain paragraph (may span multiple lines until blank)
        para_lines = [stripped]
        i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].lstrip().startswith(("#", "|", "- ", "* ", "```")):
            para_lines.append(lines[i].strip())
            i += 1
        p = doc.add_paragraph()
        add_inline_runs(p, " ".join(para_lines))


def main() -> None:
    text = MD.read_text(encoding="utf-8")
    doc = Document()

    # Tighter defaults
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    convert(text, doc)

    doc.save(DOCX)
    print(f"Wrote {DOCX} ({DOCX.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
